# import base64
# import json
# import logging
# import re
# import threading
# from concurrent.futures import ThreadPoolExecutor, as_completed
# from dataclasses import dataclass
# from io import BytesIO
# from datetime import datetime, timedelta
# from typing import Iterable, List, Optional, Tuple

# import google.generativeai as genai  # kept for compatibility if API key provided
# from google.cloud import storage
# from google.api_core.retry import Retry
# from google.cloud import aiplatform
# from vertexai import generative_models as vertex_models
# import msal
# import requests
# from django.conf import settings
# from django.db import close_old_connections
# from django.utils import timezone
# from pypdf import PdfReader
# from docx import Document

# from django.core.mail import send_mail

# from .models import ErrorLog, JobDescription, Resume, ResumeScore, ResumeSource

# logger = logging.getLogger(__name__)


# def record_error(location: str, message: str, details: str = "", context: Optional[dict] = None, resume: Optional[Resume] = None) -> None:
#     try:
#         ErrorLog.objects.create(
#             location=location,
#             message=truncate(message, 500) if "truncate" in globals() else message[:500],
#             details=truncate(details, 2000) if "truncate" in globals() else details[:2000],
#             context=context or {},
#             resume=resume,
#         )
#     except Exception:
#         logger.exception("Failed to persist error log", extra={"location": location})


# def extract_text_from_pdf(content: bytes) -> str:
#     reader = PdfReader(BytesIO(content))
#     text_parts: List[str] = []
#     for page in reader.pages:
#         text_parts.append(page.extract_text() or "")
#     return "\n".join(text_parts)


# def extract_text_from_docx(content: bytes) -> str:
#     document = Document(BytesIO(content))
#     parts = [paragraph.text for paragraph in document.paragraphs]
#     return "\n".join(parts)


# def extract_text_from_attachment(name: str, content_bytes: bytes) -> str:
#     if name.lower().endswith(".pdf"):
#         return extract_text_from_pdf(content_bytes)
#     if name.lower().endswith((".docx", ".doc")):
#         try:
#             return extract_text_from_docx(content_bytes)
#         except Exception:
#             # Fall through to text decode if the DOC/DOCX parsing fails
#             pass
#     try:
#         return content_bytes.decode("utf-8", errors="ignore")
#     except Exception:
#         return ""


# def clean_string(value: str) -> str:
#     # Remove NULs and trim whitespace to avoid DB issues
#     return value.replace("\x00", "").strip()


# def truncate(value: str, length: int = 2000) -> str:
#     if not value:
#         return value
#     return value[:length]


# def extract_email_from_text(text: str) -> Optional[str]:
#     match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
#     return match.group(0) if match else None


# class GeminiEvaluator:
#     def __init__(self, api_key: Optional[str] = None):
#         self.api_key = api_key or settings.GEMINI_API_KEY
#         self.vertex_project = settings.VERTEX_PROJECT_ID
#         self.vertex_location = settings.VERTEX_LOCATION
#         self.vertex_model_name = settings.VERTEX_MODEL or "gemini-1.5-pro"
#         self.model = None

#         if self.api_key:
#             genai.configure(api_key=self.api_key)
#             # self.model = genai.GenerativeModel("gemini-2.5-pro")
#             self.model = genai.GenerativeModel("gemini-2.5-flash-lite")
#             logger.info("Using Gemini API key authentication.")
#         elif self.vertex_project and self.vertex_location:
#             try:
#                 aiplatform.init(project=self.vertex_project, location=self.vertex_location)
#                 self.model = vertex_models.GenerativeModel(self.vertex_model_name)
#                 logger.info(
#                     "Using Vertex AI service account auth.",
#                     extra={"project": self.vertex_project, "location": self.vertex_location, "model": self.vertex_model_name},
#                 )
#             except Exception as exc:
#                 logger.exception("Vertex AI init failed; falling back to keyword scoring", extra={"error": str(exc)})
#         else:
#             logger.warning("No Gemini API key or Vertex config; using fallback keyword scoring.")

#     def _extract_text(self, response) -> str:
#         # Try to pull raw text from response or first candidate parts
#         if hasattr(response, "text") and response.text:
#             return response.text
#         try:
#             parts = []
#             candidate = response.candidates[0]
#             content = getattr(candidate, "content", None)
#             if content and hasattr(content, "parts"):
#                 for part in content.parts:
#                     if hasattr(part, "text") and part.text:
#                         parts.append(part.text)
#             return "\n".join(parts)
#         except Exception:
#             return ""

#     def _parse_json(self, raw: str) -> dict:
#         if not raw:
#             return {}
#         raw = raw.strip()
#         # Remove code fences if present
#         if raw.startswith("```"):
#             raw = raw.strip("`")
#             if raw.lower().startswith("json"):
#                 raw = raw[4:]
#         try:
#             return json.loads(raw)
#         except Exception:
#             # Try to extract the first JSON object
#             import re
#             match = re.search(r"\{.*\}", raw, re.DOTALL)
#             if match:
#                 try:
#                     return json.loads(match.group(0))
#                 except Exception:
#                     return {}
#         return {}

#     def score_resume_against_job(self, resume_text: str, job: JobDescription) -> Tuple[float, list]:
#         criteria = list(job.criteria.all())
#         if not criteria:
#             return 0.0, []

#         if not self.model:
#             return self._keyword_score(resume_text, criteria)

#         prompt = (
#             "You are an HR specialist screening a resume. Do not rewrite or alter the provided criteria text. "
#             "For each criterion, return a score 0-100 based only on evidence in the resume and a brief note if the score is low. "
#             "Compute total_score as the average of all criterion scores. Respond with JSON ONLY (no prose, no code fences) "
#             "in this shape: {\"total_score\": number, \"criteria\": [{\"criterion_id\": id, \"title\": \"criterion text exactly as provided\", \"score\": number, \"notes\": \"short why\"}]}"
#         )
#         criteria_payload = [{"id": c.id, "detail": c.detail} for c in criteria]
#         clipped_resume = resume_text[:6000]
#         content = [
#             prompt,
#             f"Job: {job.name}\nSummary: {job.summary}\nCriteria: {json.dumps(criteria_payload)}\nResume:\n{clipped_resume}",
#         ]
#         try:
#             logger.info("Gemini scoring start", extra={"job": job.id, "resume_len": len(clipped_resume)})
#             response = self.model.generate_content(content) if hasattr(self.model, "generate_content") else self.model.generate_content(
#                 contents=content
#             )
#             raw = self._extract_text(response)
#             parsed = self._parse_json(raw)
#             detail = parsed.get("criteria", [])
#             total = float(parsed.get("total_score", 0))
#             logger.info(
#                 "Gemini scoring success",
#                 extra={"job": job.id, "criteria_count": len(detail), "total_score": total},
#             )
#             return total, detail
#         except Exception as exc:
#             logger.exception("Gemini scoring failed, falling back", extra={"job": job.id, "error": str(exc)})
#             return self._keyword_score(resume_text, criteria)

#     def _keyword_score(self, resume_text: str, criteria: Iterable) -> Tuple[float, list]:
#         resume_lower = resume_text.lower()
#         detail = []
#         total = 0.0
#         for c in criteria:
#             tokens = re.split(r"[,\n;/]+", c.detail)
#             tokens = [t.strip().lower() for t in tokens if len(t.strip()) > 3]
#             token_hits = sum(1 for t in tokens if t and t in resume_lower)
#             score = 0.0
#             if tokens:
#                 score = (token_hits / len(tokens)) * 100.0
#             detail.append(
#                 {
#                     "criterion_id": c.id,
#                     "title": c.detail,
#                     "score": score,
#                     "notes": "Fallback keyword match",
#                 }
#             )
#             total += score
#         if criteria:
#             total = total / len(criteria)
#         return total, detail


# @dataclass
# class AttachmentPayload:
#     message_id: str
#     sender: Optional[str]
#     attachment_name: str
#     received_at: timezone.datetime
#     content: bytes
#     source: str = ResumeSource.EMAIL


# def extract_phone_number(text: str) -> str:
#     # Basic phone matcher allowing +, digits, spaces, dashes, parentheses
#     match = re.search(r"\+?\d[\d\s\-\(\)]{6,}", text)
#     return match.group(0).strip() if match else ""


# def infer_active_seeker(sender: Optional[str], candidate_email: Optional[str], source: str) -> bool:
#     sender_l = (sender or "").lower()
#     email_l = (candidate_email or "").lower()
#     if source == ResumeSource.UPLOAD:
#         return True
#     if "linkedin" in sender_l:
#         return True
#     if sender_l and email_l and sender_l == email_l:
#         return True
#     return False


# class GoogleStorageUploader:
#     def __init__(self, bucket_name: Optional[str] = None):
#         self.bucket_name = bucket_name or settings.GCS_BUCKET_NAME
#         self.client = None
#         if self.bucket_name:
#             self.client = storage.Client()

#     def upload(self, name: str, content: bytes, content_type: Optional[str] = None) -> Optional[str]:
#         if not self.client or not self.bucket_name:
#             return None
#         bucket = self.client.bucket(self.bucket_name)
#         blob = bucket.blob(name)
#         try:
#             blob.upload_from_string(content, content_type=content_type, timeout=120, retry=Retry(deadline=180))
#             # Store the blob path; signed URLs are generated on demand for display/download
#             return blob.name
#         except Exception as exc:
#             logger.exception("GCS upload failed", extra={"blob": name, "error": str(exc)})
#             record_error("gcs_upload", "GCS upload failed", details=str(exc), context={"blob": name})
#             return None

#     def signed_url_for_blob(self, blob_name: str) -> Optional[str]:
#         if not blob_name:
#             return None
#         if not self.bucket_name:
#             return None
#         # If client unavailable (e.g., missing credentials), return a best-effort public URL
#         if not self.client:
#             return f"https://storage.googleapis.com/{self.bucket_name}/{blob_name}"
#         bucket = self.client.bucket(self.bucket_name)
#         blob = bucket.blob(blob_name)
#         try:
#             return blob.generate_signed_url(
#                 version="v4",
#                 expiration=timedelta(seconds=getattr(settings, "GCS_SIGNED_URL_TTL_SECONDS", 604800)),
#                 method="GET",
#                 response_disposition="inline",
#             )
#         except Exception:
#             return blob.public_url or f"https://storage.googleapis.com/{self.bucket_name}/{blob_name}"

#     def download_blob(self, blob_name: str) -> Optional[bytes]:
#         if not self.client or not self.bucket_name or not blob_name:
#             return None
#         try:
#             bucket = self.client.bucket(self.bucket_name)
#             blob = bucket.blob(blob_name)
#             return blob.download_as_bytes()
#         except Exception as exc:
#             logger.exception("Blob download failed", extra={"blob": blob_name, "error": str(exc)})
#             return None


# class M365InboxReader:
#     def __init__(self, config=None):
#         self.config = config or settings.M365_CONFIG

#     def is_configured(self, log: bool = True) -> bool:
#         required = ["tenant_id", "client_id", "client_secret"]
#         missing = [key for key in required if not self.config.get(key)]
#         if missing and log:
#             logger.warning("M365 configuration missing keys", extra={"missing_keys": missing})
#             return False
#         return True

#     def _client(self):
#         if not self.is_configured(log=False):
#             return None
#         authority = f"https://login.microsoftonline.com/{self.config.get('tenant_id')}"
#         return msal.ConfidentialClientApplication(
#             self.config.get("client_id"),
#             authority=authority,
#             client_credential=self.config.get("client_secret"),
#         )

#     def _fetch_attachments_for_message(self, message: dict, headers: dict, user: str) -> List[AttachmentPayload]:
#         """Fetch attachments for a single message; intended for thread pool use."""
#         message_id = message["id"]
#         sender = (message.get("from") or {}).get("emailAddress", {}).get("address")
#         received_at = timezone.make_aware(datetime.strptime(message["receivedDateTime"], "%Y-%m-%dT%H:%M:%SZ"))
#         attach_url = f"https://graph.microsoft.com/v1.0/users/{user}/messages/{message_id}/attachments"
#         try:
#             attach_raw = requests.get(attach_url, headers=headers, timeout=30)
#             attach_resp = attach_raw.json()
#             if attach_raw.status_code >= 400 or attach_resp.get("error"):
#                 logger.error(
#                     "Attachment fetch failed",
#                     extra={
#                         "message_id": message_id,
#                         "status": attach_raw.status_code,
#                         "error": attach_resp.get("error"),
#                     },
#                 )
#                 return []
#         except Exception as exc:
#             logger.exception("Attachment fetch failed", extra={"message_id": message_id, "error": str(exc)})
#             return []

#         found: List[AttachmentPayload] = []
#         for attachment in attach_resp.get("value", []):
#             if not attachment.get("contentBytes"):
#                 continue
#             name = attachment.get("name", "attachment")
#             if not name.lower().endswith((".pdf", ".txt", ".docx", ".doc")):
#                 continue
#             try:
#                 content = base64.b64decode(attachment["contentBytes"])
#             except Exception:
#                 continue
#             found.append(
#                 AttachmentPayload(
#                     message_id=message_id,
#                     sender=sender,
#                     attachment_name=name,
#                     received_at=received_at,
#                     content=content,
#                     source=ResumeSource.EMAIL,
#                 )
#             )
#         return found

#     def fetch_recent_attachments(self, limit: int = 5, fetch_all: bool = False) -> List[AttachmentPayload]:
#         client = self._client()
#         if not client:
#             return []

#         token = client.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
#         if "access_token" not in token:
#             logger.error(
#                 "M365 token acquisition failed",
#                 extra={
#                     "error": token.get("error"),
#                     "error_description": token.get("error_description"),
#                     "correlation_id": token.get("correlation_id"),
#                 },
#             )
#             return []

#         headers = {"Authorization": f"Bearer {token['access_token']}"}
#         user = self.config.get("user_id", "me")
#         top = min(limit, 100) if limit else 50
#         base_url = f"https://graph.microsoft.com/v1.0/users/{user}/messages"
#         # Primary query filters to attachments; fallback retries without filter if the mailbox rejects it as inefficient
#         messages_url = f"{base_url}?$filter=hasAttachments eq true&$orderby=receivedDateTime desc&$top={top}"
#         fallback_url = f"{base_url}?$orderby=receivedDateTime desc&$top={top}"
#         messages = []
#         next_url = messages_url
#         attempted_fallback = False
#         while next_url:
#             resp_raw = requests.get(next_url, headers=headers, timeout=30)
#             resp = resp_raw.json()
#             if resp_raw.status_code >= 400 or resp.get("error"):
#                 error_code = (resp.get("error") or {}).get("code")
#                 if not attempted_fallback and error_code == "InefficientFilter":
#                     logger.warning(
#                         "M365 message fetch inefficient filter; retrying without hasAttachments filter",
#                         extra={"url": next_url},
#                     )
#                     next_url = fallback_url
#                     attempted_fallback = True
#                     continue
#                 logger.error(
#                     "M365 message fetch failed",
#                     extra={
#                         "status": resp_raw.status_code,
#                         "error": resp.get("error"),
#                         "url": next_url,
#                     },
#                 )
#                 break
#             messages.extend(resp.get("value", []))
#             if not fetch_all and len(messages) >= limit:
#                 messages = messages[:limit]
#                 break
#             next_url = resp.get("@odata.nextLink") if fetch_all else None

#         attachments: List[AttachmentPayload] = []
#         if not messages:
#             return attachments

#         max_workers = min(8, len(messages))
#         with ThreadPoolExecutor(max_workers=max_workers) as executor:
#             future_map = {
#                 executor.submit(self._fetch_attachments_for_message, message, headers, user): message.get("id")
#                 for message in messages
#             }
#             for future in as_completed(future_map):
#                 try:
#                     attachments.extend(future.result() or [])
#                 except Exception as exc:
#                     logger.exception("Attachment worker failed", extra={"message_id": future_map[future], "error": str(exc)})
#         return attachments


# class ResumeIngestor:
#     def __init__(
#         self,
#         scorer: Optional[GeminiEvaluator] = None,
#         uploader: Optional[GoogleStorageUploader] = None,
#         max_workers: Optional[int] = None,
#         score_workers: Optional[int] = None,
#     ):
#         self.scorer = scorer  # Lazily resolved per thread to avoid cross-thread client reuse issues
#         self.uploader = uploader or GoogleStorageUploader()
#         self.max_workers = max_workers or getattr(settings, "INGEST_MAX_WORKERS", 4)
#         self.score_workers = score_workers or getattr(settings, "INGEST_SCORE_WORKERS", self.max_workers)
#         self._scorer_local = threading.local()

#     def _get_scorer(self) -> GeminiEvaluator:
#         """Thread-local scorer to keep Gemini client usage safe under concurrency."""
#         if not getattr(self._scorer_local, "instance", None):
#             self._scorer_local.instance = self.scorer or GeminiEvaluator()
#         return self._scorer_local.instance

#     # def _score_resume_for_jobs(self, resume: Resume, text_content: str, jobs: List[JobDescription]) -> None:
#     #     if not jobs:
#     #         return
#     #     max_workers = max(1, min(self.score_workers, len(jobs)))

#     #     def score_job(job: JobDescription):
#     #         close_old_connections()
#     #         scorer = self._get_scorer()
#     #         try:
#     #             total, detail = scorer.score_resume_against_job(text_content, job)
#     #             return job, total, detail
#     #         finally:
#     #             close_old_connections()

#     #     results = []
#     #     with ThreadPoolExecutor(max_workers=max_workers) as executor:
#     #         future_map = {executor.submit(score_job, job): job for job in jobs}
#     #         for future in as_completed(future_map):
#     #             try:
#     #                 job, total, detail = future.result()
#     #                 results.append((job, total, detail))
#     #             except Exception as exc:
#     #                 logger.exception("Scoring failed", extra={"job_id": getattr(future_map[future], "id", None), "error": str(exc)})
#     #                 record_error(
#     #                     location="ingest:score_job",
#     #                     message="Resume scoring failed",
#     #                     details=str(exc),
#     #                     context={"job_id": getattr(future_map[future], "id", None), "resume_id": getattr(resume, "id", None)},
#     #                 )

#     def _score_resume_for_jobs(self, resume: Resume, text_content: str, jobs: List[JobDescription]) -> None:
#         if not jobs:
#             return
#         max_workers = max(1, min(self.score_workers, len(jobs)))

#         def score_job(job: JobDescription):
#             close_old_connections()
#             scorer = self._get_scorer()
#             try:
#                 total, detail = scorer.score_resume_against_job(text_content, job)
#                 return job, total, detail
#             finally:
#                 close_old_connections()

#         results = []
#         with ThreadPoolExecutor(max_workers=max_workers) as executor:
#             future_map = {executor.submit(score_job, job): job for job in jobs}
#             for future in as_completed(future_map):
#                 try:
#                     job, total, detail = future.result()
#                     results.append((job, total, detail))
#                 except Exception as exc:
#                     logger.exception("Scoring failed", extra={"job_id": getattr(future_map[future], "id", None), "error": str(exc)})
#                     record_error(
#                         location="ingest:score_job",
#                         message="Resume scoring failed",
#                         details=str(exc),
#                         context={"job_id": getattr(future_map[future], "id", None), "resume_id": getattr(resume, "id", None)},
#                     )

#         for job, total, detail in results:
#             # 1. Create/Update the score object
#             score_obj, created = ResumeScore.objects.update_or_create(
#                 resume=resume,
#                 job=job,
#                 defaults={"total_score": total, "detail": detail},
#             )
            
#             # 2. TRIGGER AUTOMATION (Auto-Reject / Shortlist)
#             # This calls the function we added to the bottom of the file earlier
#             try:
#                 check_and_process_automation(score_obj)
#             except Exception as e:
#                 logger.error(f"Automation trigger failed for resume {resume.id}: {e}")

#         for job, total, detail in results:
#             ResumeScore.objects.update_or_create(
#                 resume=resume,
#                 job=job,
#                 defaults={"total_score": total, "detail": detail},
#             )

#     def _record_failure(self, attachment: AttachmentPayload, jobs: List[JobDescription], error: Exception) -> None:
#         """Persist a failure placeholder so it shows on the dashboard."""
#         try:
#             sender_clean = clean_string(attachment.sender or "") if attachment.sender else None
#             candidate_name = clean_string(attachment.attachment_name.rsplit(".", 1)[0])[:255]
#             resume, _ = Resume.objects.get_or_create(
#                 message_id=attachment.message_id,
#                 attachment_name=attachment.attachment_name,
#                 defaults={
#                     "sender": sender_clean,
#                     "candidate_email": sender_clean,
#                     "candidate_phone": "",
#                     "candidate_name": candidate_name,
#                     "text_content": "",
#                     "file_url": "",
#                     "received_at": attachment.received_at,
#                     "source": attachment.source or ResumeSource.EMAIL,
#                     "is_active_seeker": infer_active_seeker(sender_clean, sender_clean, attachment.source),
#                 },
#             )
#             detail = [{"title": "Processing failed", "score": 0, "notes": truncate(str(error), 200)}]
#             for job in jobs:
#                 ResumeScore.objects.update_or_create(
#                     resume=resume,
#                     job=job,
#                     defaults={"total_score": 0.0, "detail": detail},
#                 )
#             record_error(
#                 location="ingest:_record_failure",
#                 message="Resume ingest failed",
#                 details=str(error),
#                 context={"message_id": attachment.message_id, "attachment": attachment.attachment_name},
#                 resume=resume,
#             )
#         except Exception as exc:
#             logger.exception(
#                 "Failed to record ingest failure",
#                 extra={"message_id": attachment.message_id, "attachment": attachment.attachment_name, "error": str(exc)},
#             )

#     def _ingest_single(self, attachment: AttachmentPayload, jobs: List[JobDescription]) -> int:
#         close_old_connections()
#         try:
#             raw_text = extract_text_from_attachment(attachment.attachment_name, attachment.content)
#             text_content = clean_string(raw_text)
#             phone = extract_phone_number(text_content)
#             text_email = extract_email_from_text(text_content)
#             storage_path = f"resumes/{attachment.message_id}/{attachment.attachment_name}"
#             content_type = "application/pdf" if attachment.attachment_name.lower().endswith(".pdf") else "text/plain"
#             file_url = truncate(
#                 self.uploader.upload(storage_path, attachment.content, content_type=content_type) or "",
#                 2000,
#             )
#             candidate_name = clean_string(attachment.attachment_name.rsplit(".", 1)[0])[:255]
#             sender_clean = clean_string(attachment.sender or "") if attachment.sender else None
#             candidate_email = text_email or sender_clean
#             resume, created = Resume.objects.get_or_create(
#                 message_id=attachment.message_id,
#                 attachment_name=attachment.attachment_name,
#                 defaults={
#                     "sender": sender_clean,
#                     "candidate_email": candidate_email,
#                     "candidate_phone": phone,
#                     "candidate_name": candidate_name,
#                     "text_content": text_content,
#                     "file_url": file_url or "",
#                     "received_at": attachment.received_at,
#                     "source": attachment.source or ResumeSource.EMAIL,
#                     "is_active_seeker": infer_active_seeker(sender_clean, candidate_email, attachment.source),
#                 },
#             )
#             if not created:
#                 # If a previous attempt failed (empty text_content), refresh data; otherwise respect dedup
#                 if resume.text_content:
#                     return 0
#                 resume.sender = sender_clean
#                 resume.candidate_email = candidate_email
#                 resume.candidate_phone = phone
#                 resume.candidate_name = candidate_name
#                 resume.text_content = text_content
#                 resume.file_url = file_url or ""
#                 resume.received_at = attachment.received_at
#                 resume.source = attachment.source or ResumeSource.EMAIL
#                 resume.is_active_seeker = infer_active_seeker(sender_clean, candidate_email, attachment.source)
#                 resume.save(update_fields=["sender", "candidate_email", "candidate_phone", "candidate_name", "text_content", "file_url", "received_at", "source", "is_active_seeker"])

#             self._score_resume_for_jobs(resume, text_content, jobs)
#             return 1
#         except Exception as exc:
#             logger.exception(
#                 "Resume ingest failed",
#                 extra={"message_id": attachment.message_id, "attachment": attachment.attachment_name, "error": str(exc)},
#             )
#             self._record_failure(attachment, jobs, exc)
#             record_error(
#                 location="ingest:_ingest_single",
#                 message="Resume ingest failed",
#                 details=str(exc),
#                 context={"message_id": attachment.message_id, "attachment": attachment.attachment_name},
#             )
#             return 0
#         finally:
#             close_old_connections()

#     def ingest(self, attachments: List[AttachmentPayload], jobs: Optional[List[JobDescription]] = None) -> int:
#         if not attachments:
#             return 0

#         jobs = jobs or list(JobDescription.objects.filter(active=True))
#         max_workers = max(1, min(self.max_workers, len(attachments)))
#         saved_count = 0
#         with ThreadPoolExecutor(max_workers=max_workers) as executor:
#             futures = [executor.submit(self._ingest_single, attachment, jobs) for attachment in attachments]
#             for future in as_completed(futures):
#                 try:
#                     saved_count += future.result() or 0
#                 except Exception as exc:
#                     logger.exception("Ingest worker failed", extra={"error": str(exc)})
#         return saved_count


# def populate_resume_text(resume: Resume, uploader: Optional[GoogleStorageUploader] = None) -> bool:
#     """Attempt to hydrate missing resume text from stored files."""
#     if resume.text_content:
#         return False
#     uploader = uploader or GoogleStorageUploader()
#     content_bytes: Optional[bytes] = None
#     try:
#         if resume.file_url:
#             if resume.file_url.startswith("http"):
#                 resp = requests.get(resume.file_url, timeout=20)
#                 if resp.status_code < 400:
#                     content_bytes = resp.content
#             else:
#                 content_bytes = uploader.download_blob(resume.file_url)
#     except Exception as exc:
#         logger.exception("Resume text recovery failed", extra={"resume_id": resume.id, "error": str(exc)})
#         record_error(
#             location="populate_resume_text",
#             message="Resume text recovery failed",
#             details=str(exc),
#             context={"resume_id": resume.id},
#             resume=resume,
#         )
#         return False

#     if not content_bytes:
#         return False
#     text_content = clean_string(extract_text_from_attachment(resume.attachment_name, content_bytes))
#     if not text_content:
#         return False
#     resume.text_content = text_content
#     resume.save(update_fields=["text_content"])
#     return True


# def check_and_process_automation(resume_score):
#     """
#     Checks the score against the threshold.
#     - If Score < Threshold: Auto-reject and send email.
#     - If Score >= Threshold: Mark as Shortlisted.
#     """
#     resume = resume_score.resume
#     score_val = resume_score.total_score
#     # Default to 60 if not set in settings
#     threshold = getattr(settings, 'AUTO_REJECTION_THRESHOLD', 60) 

#     # 1. AUTO REJECTION LOGIC
#     if score_val < threshold:
#         # Update Status
#         resume.status = 'AUTO_REJECTED'
#         resume.save()

#         # Send Email (Only if not already sent)
#         if not resume.rejection_email_sent and resume.candidate_email:
#             subject = f"Update regarding your application at {resume_score.job.name}"
#             message = (
#                 f"Dear {resume.candidate_name or 'Candidate'},\n\n"
#                 f"Thank you for your interest in the {resume_score.job.name} position. "
#                 "We appreciate the time you took to apply.\n\n"
#                 "After reviewing your application, we have decided not to proceed with your candidacy "
#                 "at this time. We will keep your resume on file for future openings.\n\n"
#                 "Best regards,\n"
#                 "HR Team"
#             )
            
#             try:
#                 send_mail(
#                     subject,
#                     message,
#                     settings.DEFAULT_FROM_EMAIL,
#                     [resume.candidate_email],
#                     fail_silently=False,
#                 )
#                 # Mark flag as True so we don't spam them later
#                 resume.rejection_email_sent = True
#                 resume.save()
#                 print(f"✓ Rejection email sent to {resume.candidate_email}")
#             except Exception as e:
#                 print(f"✗ Failed to send email: {e}")
#                 record_error("automation:email", f"Failed to send rejection email to {resume.candidate_email}", str(e))

#     # 2. SHORTLIST LOGIC
#     else:
#         # Ensure status is NEW so it appears in the Inbox.
#         # If it was 'PENDING' during processing, move it to 'NEW'.
#         if resume.status != 'NEW':
#             resume.status = 'NEW'
#             resume.save()










import base64
import json
import logging
import re
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from io import BytesIO
from datetime import datetime, timedelta
from typing import Iterable, List, Optional, Tuple, Any

import google.generativeai as genai
from google.cloud import storage
from google.api_core.retry import Retry
from google.cloud import aiplatform
from vertexai import generative_models as vertex_models
import msal
import requests
from django.conf import settings
from django.db import close_old_connections
from django.db.models import Q  # <--- NEW IMPORT
from django.utils import timezone
from pypdf import PdfReader
from docx import Document
from.utils import log_activity
import subprocess  # <--- Add this


# Google AI Imports
from google.cloud import aiplatform
from vertexai.preview.generative_models import GenerativeModel as VertexModel # Adjust import based on your SDK version
import vertexai.generative_models as vertex_models # Standard Vertex SDK

from django.conf import settings


from django.core.mail import send_mail

from .models import ErrorLog, JobDescription, Resume, ResumeScore, ResumeSource

logger = logging.getLogger(__name__)


# def record_error(location: str, message: str, details: str = "", context: Optional[dict] = None, resume: Optional[Resume] = None) -> None:
#     try:
#         ErrorLog.objects.create(
#             location=location,
#             message=truncate(message, 500) if "truncate" in globals() else message[:500],
#             details=truncate(details, 2000) if "truncate" in globals() else details[:2000],
#             context=context or {},
#             resume=resume,
#         )
#     except Exception:
#         logger.exception("Failed to persist error log", extra={"location": location})

def record_error(location: str, message: str, details: str = "", context: Optional[dict] = None, resume: Any = None) -> None:
    """Safe error logging that won't crash if DB is down."""
    try:
        # Import inside function to avoid circular imports if necessary
        from .models import ErrorLog 
        ErrorLog.objects.create(
            location=location,
            message=truncate(message, 500),
            details=truncate(details, 2000),
            context=context or {},
            resume=resume,
        )
    except Exception as e:
        logger.error(f"FALLBACK LOG - {location}: {message} | {details} | Error: {e}")


# def extract_text_from_pdf(content: bytes) -> str:
#     reader = PdfReader(BytesIO(content))
#     text_parts: List[str] = []
#     for page in reader.pages:
#         text_parts.append(page.extract_text() or "")
#     return "\n".join(text_parts)

def extract_text_from_pdf(content: bytes, filename: str = "") -> str:
    """Extracts text from PDF. Returns '' if scanned image."""
    try:
        reader = PdfReader(BytesIO(content))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        extracted_text = "\n".join(text_parts)
        
        # Validation: If text is suspiciously short, it's likely a scan
        if len(extracted_text.strip()) < 50:
            record_error("extract_text_from_pdf", "PDF Text < 50 chars (Likely Scanned Image)", f"File: {filename}")
            return "" 
            
        return extracted_text
    except Exception as e:
        record_error("extract_text_from_pdf", f"PDF parsing failed: {str(e)}", f"File: {filename}")
        return ""


# def extract_text_from_docx(content: bytes) -> str:
#     document = Document(BytesIO(content))
#     parts = [paragraph.text for paragraph in document.paragraphs]
#     return "\n".join(parts)

def extract_text_from_docx(content: bytes, filename: str = "") -> str:
    """Extracts text from paragraphs AND tables in DOCX."""
    try:
        document = Document(BytesIO(content))
        text_parts = []

        # 1. Paragraphs
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # 2. Tables (Crucial for Skills sections)
        for table in document.tables:
            for row in table.rows:
                # Join cells with a pipe | to keep structure
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        text = "\n".join(text_parts)
        if not text:
            record_error("extract_text_from_docx", "Empty DOCX extraction", f"File: {filename}")
        return text

    except Exception as e:
        record_error("extract_text_from_docx", f"DOCX parsing failed: {str(e)}", f"File: {filename}")
        return ""


# def extract_text_from_attachment(name: str, content_bytes: bytes) -> str:
#     if name.lower().endswith(".pdf"):
#         return extract_text_from_pdf(content_bytes)
#     if name.lower().endswith((".docx", ".doc")):
#         try:
#             return extract_text_from_docx(content_bytes)
#         except Exception:
#             pass
#     try:
#         return content_bytes.decode("utf-8", errors="ignore")
#     except Exception:
#         return ""

def extract_text_from_attachment(name: str, content_bytes: bytes) -> str:
    """Master routing function."""
    name_lower = name.lower()
    text = ""
    
    if name_lower.endswith(".pdf"):
        text = extract_text_from_pdf(content_bytes, name)
    elif name_lower.endswith(".docx"):
        text = extract_text_from_docx(content_bytes, name)
    elif name_lower.endswith(".doc"):
        record_error("extract_text_from_attachment", "Unsupported .doc format", f"File: {name}")
    else:
        # Fallback
        try:
            text = content_bytes.decode("utf-8", errors="ignore")
        except Exception:
            pass

    return clean_string(text)


# def clean_string(value: str) -> str:
#     return value.replace("\x00", "").strip()

def clean_string(value: str) -> str:
    if not value: return ""
    return value.replace("\x00", "").strip()


# def truncate(value: str, length: int = 2000) -> str:
#     if not value:
#         return value
#     return value[:length]

def truncate(value: str, length: int = 2000) -> str:
    if not value: return ""
    return str(value)[:length]


def extract_email_from_text(text: str) -> Optional[str]:
    match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    return match.group(0) if match else None



import io
import pypdf

# def extract_text_with_links(pdf_file_path_or_stream):
#     """
#     Extracts visible text AND hidden hyperlinks from a PDF.
#     Handles bytes, file paths, or file-like objects automatically.
#     """
#     try:
#         # === FIX START: Auto-convert bytes to stream ===
#         if isinstance(pdf_file_path_or_stream, bytes):
#             pdf_file_path_or_stream = io.BytesIO(pdf_file_path_or_stream)
#         # === FIX END ===

#         reader = pypdf.PdfReader(pdf_file_path_or_stream)
#         full_text = []
#         found_links = set()

#         for page in reader.pages:
#             # 1. Extract Visible Text
#             text = page.extract_text() or ""
#             full_text.append(text)

#             # 2. Extract Hidden Hyperlinks (Metadata)
#             if "/Annots" in page:
#                 # Resolve indirect object reference if necessary
#                 annots = page["/Annots"]
#                 # Handle case where /Annots is an IndirectObject
#                 if hasattr(annots, "get_object"):
#                      annots = annots.get_object()
                
#                 for annot in annots:
#                     try:
#                         obj = annot.get_object()
#                         if "/A" in obj and "/URI" in obj["/A"]:
#                             uri = obj["/A"]["/URI"]
#                             # Only grab useful links
#                             if "linkedin.com" in uri or "github.com" in uri:
#                                 found_links.add(uri)
#                     except Exception:
#                         continue 

#         # Combine them
#         final_content = "\n".join(full_text)
        
#         if found_links:
#             final_content += "\n\n--- DETECTED HYPERLINKS ---\n"
#             final_content += "\n".join(found_links)
            
#         return final_content

#     except Exception as e:
#         print(f"PDF Link Extraction failed: {e}")
#         # Return empty string so the pipeline doesn't crash completely
#         return ""

import pypdf
import io
import docx  # Import python-docx

# def extract_text_with_links(file_content, filename=""):
#     """
#     Extracts text/links from PDF or DOCX.
#     detects format based on filename extension or file header.
#     """
#     try:
#         # 1. Normalize Input to Stream (BytesIO)
#         if hasattr(file_content, 'read') and hasattr(file_content, 'seek'):
#             stream = file_content
#             stream.seek(0)
#         else:
#             stream = io.BytesIO(file_content)

#         # 2. Detect Format (Simple check using filename or magic numbers)
#         is_pdf = False
#         if filename.lower().endswith('.pdf'):
#             is_pdf = True
#         else:
#             # Fallback: Peek at first 4 bytes for PDF header (%PDF)
#             current_pos = stream.tell()
#             header = stream.read(4)
#             stream.seek(current_pos) # Reset
#             if header == b'%PDF':
#                 is_pdf = True

#         # 3. ROUTE TO CORRECT EXTRACTOR
#         if is_pdf:
#             return _extract_pdf(stream)
#         else:
#             return _extract_docx(stream)

#     except Exception as e:
#         print(f"Extraction failed: {e}")
#         return ""

# def _extract_pdf(stream):
#     """Helper for PDF Extraction"""
#     try:
#         reader = pypdf.PdfReader(stream)
#         full_text = []
#         found_links = set()

#         for page in reader.pages:
#             text = page.extract_text() or ""
#             full_text.append(text)
            
#             if "/Annots" in page:
#                 for annot in page["/Annots"]:
#                     try:
#                         obj = annot.get_object()
#                         if "/A" in obj and "/URI" in obj["/A"]:
#                             uri = obj["/A"]["/URI"]
#                             if any(x in uri for x in ["linkedin.com", "github.com", "portfolio"]):
#                                 found_links.add(uri)
#                     except: continue

#         content = "\n".join(full_text)
#         if found_links:
#             content += "\n\n--- DETECTED HYPERLINKS ---\n" + "\n".join(found_links)
#         return content
#     except Exception as e:
#         print(f"PDF error: {e}")
#         return ""

# def _extract_docx(stream):
#     """Helper for DOCX Extraction"""
#     try:
#         doc = docx.Document(stream)
#         full_text = []
#         found_links = set()
        
#         # Extract Paragraphs
#         for para in doc.paragraphs:
#             full_text.append(para.text)
            
#             # --- Extract Hyperlinks from DOCX XML ---
#             # python-docx doesn't support links natively, we must parse XML
#             if para._p.xml:
#                 # Look for <w:hyperlink> tags in the XML
#                 for rel in para.part.rels.values():
#                     if "hyperlink" in rel.reltype:
#                         if "linkedin.com" in rel.target_ref or "github.com" in rel.target_ref:
#                             found_links.add(rel.target_ref)

#         content = "\n".join(full_text)
#         if found_links:
#             content += "\n\n--- DETECTED HYPERLINKS ---\n" + "\n".join(found_links)
#         return content
#     except Exception as e:
#         print(f"DOCX error: {e}")
#         return ""





import io
import logging
import os
import re
import tempfile
import pypdf
import docx
import pytesseract
from pdf2image import convert_from_bytes
# import textract
from zipfile import BadZipFile

# Configure Logging
logger = logging.getLogger(__name__)

# --- CONFIGURATION ---
# If Tesseract is not in your PATH, set it here (Windows users specifically)
# pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract' 

def extract_text_with_links(file_content, filename=""):
    """
    Master extraction function.
    Strategy:
    1. Try Fast Extraction (pypdf / python-docx).
    2. If text is empty or fails, trigger 'Costly Fallback' (OCR / Textract).
    """
    text = ""
    
    # 1. Normalize Input to Stream
    if isinstance(file_content, (bytes, bytearray)):
        stream = io.BytesIO(file_content)
        file_bytes = file_content
    elif hasattr(file_content, 'read'):
        file_content.seek(0)
        file_bytes = file_content.read()
        stream = io.BytesIO(file_bytes)
    else:
        logger.error("Invalid file content type")
        return ""

    # 2. Detect Format
    is_pdf = filename.lower().endswith('.pdf')
    is_docx = filename.lower().endswith('.docx') or filename.lower().endswith('.doc')

    if not is_pdf and not is_docx:
        # Magic Number Check
        if file_bytes.startswith(b'%PDF'):
            is_pdf = True
        elif file_bytes.startswith(b'PK') or file_bytes.startswith(b'\xD0\xCF\x11\xE0'): # PK=Zip, D0CF=OLE(doc)
            is_docx = True

    # 3. FAST EXTRACTION ATTEMPT
    try:
        if is_pdf:
            text = _extract_pdf_fast(stream)
        elif is_docx:
            text = _extract_docx_fast(stream)
    except Exception as e:
        logger.warning(f"Fast extraction failed for {filename}: {e}")

    # 4. COSTLY FALLBACK (If fast failed or returned minimal text)
    # We define "failure" as having less than 50 characters of text
    if len(text.strip()) < 50:
        logger.info(f"Triggering COSTLY extraction for {filename}")
        try:
            if is_pdf:
                text = _extract_pdf_ocr(file_bytes)
            elif is_docx:
                text = _extract_doc_legacy(file_bytes, filename)
        except Exception as e:
            logger.error(f"Costly extraction also failed: {e}")

    return text

# ==========================================
# FAST METHODS (Standard Libraries)
# ==========================================

def _extract_pdf_fast(stream):
    try:
        reader = pypdf.PdfReader(stream)
        full_text = []
        found_links = set()

        for page in reader.pages:
            text = page.extract_text() or ""
            full_text.append(text)
            
            # Extract Metadata Links
            if "/Annots" in page:
                for annot in page["/Annots"]:
                    try:
                        obj = annot.get_object()
                        if "/A" in obj and "/URI" in obj["/A"]:
                            uri = obj["/A"]["/URI"]
                            if any(x in uri for x in ["linkedin.com", "github.com"]):
                                found_links.add(uri)
                    except: continue

        content = "\n".join(full_text)
        if found_links:
            content += "\n\n--- DETECTED HYPERLINKS ---\n" + "\n".join(found_links)
        return content
    except Exception:
        return ""

def _extract_docx_fast(stream):
    try:
        doc = docx.Document(stream)
        full_text = []
        found_links = set()
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            # Basic XML Link extraction
            if para._p.xml:
                try:
                    for rel in para.part.rels.values():
                        if "hyperlink" in rel.reltype:
                            if "linkedin.com" in rel.target_ref:
                                found_links.add(rel.target_ref)
                except: pass

        content = "\n".join(full_text)
        if found_links:
            content += "\n\n--- DETECTED HYPERLINKS ---\n" + "\n".join(found_links)
        return content
    except BadZipFile:
        return "" # Trigger fallback
    except Exception:
        return ""

# ==========================================
# COSTLY METHODS (OCR & Binary Parsers)
# ==========================================

def _extract_pdf_ocr(file_bytes):
    """
    Uses pdf2image to convert pages to images, then Tesseract OCR to read text.
    HEAVY PROCESS: High CPU/RAM usage.
    """
    try:
        # Convert PDF to List of Images
        images = convert_from_bytes(file_bytes)
        full_text = []
        
        for i, image in enumerate(images):
            # OCR each page
            page_text = pytesseract.image_to_string(image)
            full_text.append(page_text)
            
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"OCR Failed: {e}")
        return ""

# def _extract_doc_legacy(file_bytes, original_filename):
#     """
#     Handles .doc (Word 97-2003) and stubborn .docx files.
#     Writes to a temp file and uses 'textract' (which calls antiword/antiword).
#     """
#     try:
#         # Textract requires a file path, it can't read memory streams for .doc
#         suffix = ".doc" if not original_filename.lower().endswith(".docx") else ".docx"
        
#         with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
#             temp.write(file_bytes)
#             temp_path = temp.name

#         try:
#             # Process
#             text_bytes = textract.process(temp_path)
#             text = text_bytes.decode("utf-8", errors='ignore')
#         finally:
#             # Cleanup temp file
#             if os.path.exists(temp_path):
#                 os.remove(temp_path)
                
#         return text
#     except Exception as e:
#         logger.error(f"Legacy Doc Extraction Failed: {e}")
#         return ""

def _extract_doc_legacy(file_bytes, original_filename):
    """
    Handles .doc (Word 97-2003) files using 'antiword' directly via subprocess.
    Replaces the broken 'textract' library.
    """
    try:
        # Create a temp file
        suffix = ".doc"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
            temp.write(file_bytes)
            temp_path = temp.name

        try:
            # DIRECTLY call antiword using subprocess
            # This bypasses the need for the broken textract Python library
            result = subprocess.run(
                ['antiword', temp_path], 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE
            )
            
            if result.returncode == 0:
                text = result.stdout.decode("utf-8", errors='ignore')
                return text
            else:
                logger.error(f"Antiword failed: {result.stderr.decode()}")
                return ""
                
        finally:
            # Cleanup temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        logger.error(f"Legacy Doc Extraction Failed: {e}")
        return ""


# import json
# import logging
# import re
# from typing import Optional, Tuple, Iterable, Dict, List
# import google.generativeai as genai
# from google.cloud import aiplatform
# from vertexai.preview.generative_models import GenerativeModel as VertexModel
# import vertexai.generative_models as vertex_models
# from django.conf import settings

# logger = logging.getLogger(__name__)

# class GeminiEvaluator:
#     def __init__(self, api_key: Optional[str] = None):
#         self.api_key = api_key or getattr(settings, 'GEMINI_API_KEY', None)
#         self.vertex_project = getattr(settings, 'VERTEX_PROJECT_ID', None)
#         self.vertex_location = getattr(settings, 'VERTEX_LOCATION', 'us-central1')
#         # Use a STABLE model version
#         self.vertex_model_name = getattr(settings, 'VERTEX_MODEL', "gemini-1.5-flash") 
#         self.model = None

#         if self.api_key:
#             genai.configure(api_key=self.api_key)
#             self.model = genai.GenerativeModel("gemini-2.5-flash")
#             logger.info("Using Gemini API key authentication.")
#         elif self.vertex_project and self.vertex_location:
#             try:
#                 aiplatform.init(project=self.vertex_project, location=self.vertex_location)
#                 self.model = vertex_models.GenerativeModel(self.vertex_model_name)
#                 logger.info(f"Using Vertex AI auth with model: {self.vertex_model_name}")
#             except Exception as exc:
#                 logger.exception("Vertex AI init failed")
#         else:
#             logger.warning("No Gemini config found. Falling back to keyword matching.")

#     def _extract_text(self, response) -> str:
#         """Safe extraction from Gemini response object."""
#         try:
#             if hasattr(response, "text") and response.text:
#                 return response.text
#             # Fallback for Vertex AI objects
#             if hasattr(response, "candidates") and response.candidates:
#                 return response.candidates[0].content.parts[0].text
#         except Exception:
#             pass
#         return ""

#     def _parse_json(self, raw: str) -> dict:
#         if not raw: return {}
        
#         # Clean Markdown ```json ... ```
#         raw = re.sub(r"^```json", "", raw, flags=re.MULTILINE)
#         raw = re.sub(r"^```", "", raw, flags=re.MULTILINE)
#         raw = raw.strip()
        
#         try:
#             return json.loads(raw)
#         except json.JSONDecodeError:
#             # Try to find the first '{' and last '}'
#             match = re.search(r"\{.*\}", raw, re.DOTALL)
#             if match:
#                 try:
#                     return json.loads(match.group(0))
#                 except:
#                     pass
#             return {}

#     def score_resume_against_job(self, resume_text: str, job) -> Tuple[float, list, dict]:
#         """
#         Returns: (Total Score, Criteria Details List, Candidate Profile Dict)
#         """
#         criteria = list(job.criteria.all())
#         if not criteria:
#             return 0.0, [], {}

#         if not self.model:
#             return self._keyword_score(resume_text, criteria)

#         prompt = (
#             "You are an expert HR Technical Recruiter. Your task is to analyze a resume against a job description.\n\n"
#             "OUTPUT FORMAT:\n"
#             "Return ONLY valid JSON with this exact structure:\n"
#             "{\n"
#             "  \"candidate_profile\": {\n"
#             "    \"name\": \"Full Name or not Found\",\n"
#             "    \"email\": \"Email or not Found\",\n"
#             "    \"phone\": \"Phone or not Found\",\n"
#             "    \"linkedin_url\": \"URL or not Found\",\n"
#             "    \"current_location\": \"City, Country or not Found\",\n"
#             "    \"current_company\": \"Latest Company or not Found\",\n"
#             "    \"current_role\": \"Latest Job Title or not Found\"\n"
#             "  },\n"
#             "  \"scores\": {\n"
#             "     \"total_score\": number (0-100),\n"
#             "     \"reasoning\": \"Brief summary of fit\",\n"
#             "     \"criteria_analysis\": [\n"
#             "        {\"criterion_id\": id, \"title\": \"text\", \"score\": number, \"evidence\": \"short reason\"}\n"
#             "     ]\n"
#             "  }\n"
#             "}\n\n"
#             "INSTRUCTIONS:\n"
#             "1. Extract candidate metadata accurately.\n"
#             "2. Score each criterion 0-100 based strictly on the resume evidence.\n"
#             "3. Be strict. If a skill is missing, score it 0.\n"
#         )

#         criteria_payload = [{"id": c.id, "detail": c.detail} for c in criteria]
        
#         # Truncate to avoid context limit
#         clipped_resume = resume_text[:12000] 
        
#         content = [
#             prompt,
#             f"--- JOB: {job.name} ---\n{job.summary}\n\n"
#             f"--- CRITERIA ---\n{json.dumps(criteria_payload)}\n\n"
#             f"--- RESUME ---\n{clipped_resume}"
#         ]

#         try:
#             # Generate
#             response = self.model.generate_content(content)
#             raw_text = self._extract_text(response)
#             parsed = self._parse_json(raw_text)

#             # Unpack JSON correctly (fixing the path error)
#             profile = parsed.get("candidate_profile", {})
#             score_data = parsed.get("scores", {})
            
#             total = float(score_data.get("total_score", 0))
#             detail = score_data.get("criteria_analysis", [])
            
#             # Fallback if AI messed up the structure but returned criteria at root
#             if not detail and "criteria" in parsed:
#                 detail = parsed["criteria"]
#                 total = float(parsed.get("total_score", 0))

#             return total, detail, profile

#         except Exception as exc:
#             logger.exception("Gemini scoring failed", extra={"job": job.id, "error": str(exc)})
#             return self._keyword_score(resume_text, criteria)

#     def _keyword_score(self, resume_text: str, criteria: Iterable) -> Tuple[float, list, dict]:
#         """Fallback scorer. Returns 3 values to match signature."""
#         resume_lower = resume_text.lower()
#         detail = []
#         total = 0.0
        
#         for c in criteria:
#             # Improved token splitting
#             tokens = re.split(r"[,\n;/]+", c.detail)
#             tokens = [t.strip().lower() for t in tokens if len(t.strip()) > 2]
            
#             if not tokens:
#                 score = 0.0
#             else:
#                 # Fix duplicate counting bug: Only count 1 hit per unique token
#                 token_hits = 0
#                 for t in tokens:
#                     if t in resume_lower:
#                         token_hits += 1
                
#                 score = (token_hits / len(tokens)) * 100.0
#                 score = min(score, 100.0) # Cap at 100%
            
#             detail.append({
#                 "criterion_id": c.id,
#                 "title": c.detail,
#                 "score": score,
#                 "notes": "Fallback keyword match",
#             })
#             total += score

#         if criteria:
#             total = total / len(criteria)
            
#         # Return empty dict for profile
#         return total, detail, {}






import json
import logging
import re
from typing import Optional, Tuple, Iterable, Dict, List
import google.generativeai as genai
from google.cloud import aiplatform
from vertexai.preview.generative_models import GenerativeModel as VertexModel
import vertexai.generative_models as vertex_models
from django.conf import settings

logger = logging.getLogger(__name__)

class GeminiEvaluator:
    def __init__(self, api_key: Optional[str] = None):
        self.api_key = api_key or getattr(settings, 'GEMINI_API_KEY', None)
        self.vertex_project = getattr(settings, 'VERTEX_PROJECT_ID', None)
        self.vertex_location = getattr(settings, 'VERTEX_LOCATION', 'us-central1')
        self.vertex_model_name = getattr(settings, 'VERTEX_MODEL', "gemini-1.5-flash") 
        self.model = None

        if self.api_key:
            genai.configure(api_key=self.api_key)
            self.model = genai.GenerativeModel("gemini-2.5-flash")
            logger.info("Using Gemini API key authentication.")
        elif self.vertex_project and self.vertex_location:
            try:
                aiplatform.init(project=self.vertex_project, location=self.vertex_location)
                self.model = vertex_models.GenerativeModel(self.vertex_model_name)
                logger.info(f"Using Vertex AI auth with model: {self.vertex_model_name}")
            except Exception as exc:
                logger.exception("Vertex AI init failed")
        else:
            logger.warning("No Gemini config found. Falling back to keyword matching.")

    def _extract_text(self, response) -> str:
        """Safe extraction from Gemini response object."""
        try:
            if hasattr(response, "text") and response.text:
                return response.text
            if hasattr(response, "candidates") and response.candidates:
                return response.candidates[0].content.parts[0].text
        except Exception:
            pass
        return ""

    def _parse_json(self, raw: str) -> dict:
        if not raw: return {}
        
        # Clean Markdown ```json ... ```
        raw = re.sub(r"^```json", "", raw, flags=re.MULTILINE)
        raw = re.sub(r"^```", "", raw, flags=re.MULTILINE)
        raw = raw.strip()
        
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", raw, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except:
                    pass
            return {}

    def score_resume_against_job(self, resume_text: str, job) -> Tuple[float, list, dict]:
        """
        Returns: (Total Score 0-100, Criteria Details List (Scores 0-10), Candidate Profile Dict)
        """
        criteria = list(job.criteria.all())
        if not criteria:
            return 0.0, [], {}

        if not self.model:
            return self._keyword_score(resume_text, criteria)

        # UPDATED PROMPT: 
        # 1. Scores out of 10.
        # 2. Key 'notes' matches your HTML template.
        # 3. Explicit profile fields.
        # prompt = (
        #     "You are an expert HR Technical Recruiter. Your task is to analyze a resume against a job description.\n\n"
        #     "OUTPUT FORMAT:\n"
        #     "Return ONLY valid JSON with this exact structure:\n"
        #     "{\n"
        #     "  \"candidate_profile\": {\n"
        #     "    \"name\": \"Full Name or Not Found\",\n"
        #     "    \"email\": \"Email or Not Found\",\n"
        #     "    \"phone\": \"Phone or Not Found\",\n"
        #     "    \"linkedin_url\": \"URL or Not Found\",\n"
        #     "    \"current_location\": \"City, Country or Not Found\",\n"
        #     "    \"current_company\": \"Latest Company or Not Found\",\n"
        #     "    \"current_role\": \"Latest Job Title or Not Found\"\n"
        #     "  },\n"
        #     "  \"scores\": {\n"
        #     "     \"reasoning\": \"Brief summary of fit\",\n"
        #     "     \"criteria_analysis\": [\n"
        #     "        {\"criterion_id\": id, \"title\": \"text\", \"score\": number, \"notes\": \"Specific evidence from resume justifying the score\"}\n"
        #     "     ]\n"
        #     "  }\n"
        #     "}\n\n"
        #     "INSTRUCTIONS:\n"
        #     "1. Extract candidate metadata accurately. Return null if not found.\n"
        #     "2. Score each criterion on a scale of 0 to 10 (0=No evidence, 10=Perfect Match).\n"
        #     "3. BE STRICT. If a skill is missing, score it 0.\n"
        #     "4. The 'notes' field MUST explain the score based on the resume text.\n"
        # )

        # UPDATED PROMPT with specific URL instructions
        prompt = (
            "You are an expert HR Technical Recruiter. Your task is to analyze a resume against a job description.\n\n"
            "OUTPUT FORMAT:\n"
            "Return ONLY valid JSON with this exact structure:\n"
            "{\n"
            "  \"candidate_profile\": {\n"
            "    \"name\": null,\n"
            "    \"email\": null,\n"
            "    \"phone\": null,\n"
            "    \"linkedin_url\": null,\n"
            "    \"current_location\": null,\n"
            "    \"current_company\": null,\n"
            "    \"current_role\": null\n"
            "  },\n"
            "  \"scores\": {\n"
            "     \"reasoning\": \"Brief summary of fit\",\n"
            "     \"criteria_analysis\": [\n"
            "        {\"criterion_id\": id, \"title\": \"text\", \"score\": number, \"notes\": \"Specific evidence from resume justifying the score\"}\n"
            "     ]\n"
            "  }\n"
            "}\n\n"
            "INSTRUCTIONS:\n"
            "1. Extract candidate metadata accurately. Return 'unable to find' if not found.\n"
            "2. **FOR LINKEDIN:** Look for the longest possible URL string. If the text contains 'linkedin.com/in/name-id-123', extract the full ID. Do not truncate at hyphens.\n"
            "3. Score each criterion on a scale of 0 to 10 (0=No evidence, 10=Perfect Match).\n"
            "4. BE STRICT. If a skill is missing, score it 0.\n"
            "5. The 'notes' field MUST explain the score based on the resume text.\n"
        )

        criteria_payload = [{"id": c.id, "detail": c.detail} for c in criteria]
        clipped_resume = resume_text[:12000] 
        
        content = [
            prompt,
            f"--- JOB: {job.name} ---\n{job.summary}\n\n"
            f"--- CRITERIA ---\n{json.dumps(criteria_payload)}\n\n"
            f"--- RESUME ---\n{clipped_resume}"
        ]

        try:
            # Generate
            response = self.model.generate_content(content)
            raw_text = self._extract_text(response)
            parsed = self._parse_json(raw_text)

            # Unpack
            profile = parsed.get("candidate_profile", {})
            score_data = parsed.get("scores", {})
            detail = score_data.get("criteria_analysis", [])

            # Fallback for bad JSON structure
            if not detail and "criteria" in parsed:
                detail = parsed["criteria"]

            # --- CALCULATE TOTAL (Normalize 0-10 scale to 0-100) ---
            total = 0.0
            if detail:
                # Sum of (0-10) scores
                raw_sum = sum(float(d.get("score", 0)) for d in detail)
                max_possible = len(detail) * 10
                # Convert to percentage
                total = (raw_sum / max_possible) * 100.0

            return total, detail, profile

        except Exception as exc:
            logger.exception("Gemini scoring failed", extra={"job": job.id, "error": str(exc)})
            return self._keyword_score(resume_text, criteria)

    def _keyword_score(self, resume_text: str, criteria: Iterable) -> Tuple[float, list, dict]:
        """Fallback scorer. Returns scores out of 10 for consistency."""
        resume_lower = resume_text.lower()
        detail = []
        total_percent = 0.0
        
        for c in criteria:
            tokens = re.split(r"[,\n;/]+", c.detail)
            tokens = [t.strip().lower() for t in tokens if len(t.strip()) > 2]
            
            if not tokens:
                score_out_of_10 = 0.0
            else:
                token_hits = 0
                for t in tokens:
                    if t in resume_lower:
                        token_hits += 1
                
                # Calculate percentage then convert to 0-10
                percent = (token_hits / len(tokens)) * 100.0
                percent = min(percent, 100.0)
                score_out_of_10 = percent / 10.0 # Convert 100 -> 10
            
            detail.append({
                "criterion_id": c.id,
                "title": c.detail,
                "score": round(score_out_of_10, 1), # e.g. 7.5
                "notes": "Fallback keyword match",
            })
            total_percent += (score_out_of_10 * 10) # Summing up percentages effectively

        if criteria:
            total = total_percent / len(criteria)
        else:
            total = 0.0
            
        return total, detail, {}


# @dataclass
# class AttachmentPayload:
#     message_id: str
#     sender: Optional[str]
#     attachment_name: str
#     received_at: timezone.datetime
#     content: bytes
#     source: str = ResumeSource.EMAIL

@dataclass
class AttachmentPayload:
    message_id: str
    sender: Optional[str]
    attachment_name: str
    received_at: timezone.datetime
    content: bytes
    source: str = ResumeSource.EMAIL


# def extract_phone_number(text: str) -> str:
#     match = re.search(r"\+?\d[\d\s\-\(\)]{6,}", text)
#     return match.group(0).strip() if match else ""

def extract_phone_number(text: str) -> str:
    """
    Fallback extraction using Regex. 
    Gemini is preferred, but this catches obvious patterns if Gemini fails.
    """
    if not text:
        return ""
        
    # Improved Regex:
    # 1. Matches international codes (+91, +1)
    # 2. Avoids matching dates (YYYY-MM-DD) by ensuring no surrounding - or / checks 
    # 3. Looks for 10-15 digits allowing for spaces, dashes, brackets
    phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4,6}'
    
    matches = re.findall(phone_pattern, text)
    
    for match in matches:
        # Filter out likely dates (e.g., 2023-10-12)
        clean_num = re.sub(r'\D', '', match)
        if 10 <= len(clean_num) <= 15:
            return match.strip()
            
    return ""


def infer_active_seeker(sender: Optional[str], candidate_email: Optional[str], source: str) -> bool:
    sender_l = (sender or "").lower()
    email_l = (candidate_email or "").lower()
    if source == ResumeSource.UPLOAD:
        return True
    if "linkedin" in sender_l:
        return True
    if sender_l and email_l and sender_l == email_l:
        return True
    return False


# class GoogleStorageUploader:
#     def __init__(self, bucket_name: Optional[str] = None):
#         self.bucket_name = bucket_name or settings.GCS_BUCKET_NAME
#         self.client = None
#         if self.bucket_name:
#             self.client = storage.Client()

#     def upload(self, name: str, content: bytes, content_type: Optional[str] = None) -> Optional[str]:
#         if not self.client or not self.bucket_name:
#             return None
#         bucket = self.client.bucket(self.bucket_name)
#         blob = bucket.blob(name)
#         try:
#             blob.upload_from_string(content, content_type=content_type, timeout=120, retry=Retry(deadline=180))
#             return blob.name
#         except Exception as exc:
#             logger.exception("GCS upload failed", extra={"blob": name, "error": str(exc)})
#             return None

#     def signed_url_for_blob(self, blob_name: str) -> Optional[str]:
#         if not blob_name:
#             return None
#         if not self.bucket_name:
#             return None
#         if not self.client:
#             return f"[https://storage.googleapis.com/](https://storage.googleapis.com/){self.bucket_name}/{blob_name}"
#         bucket = self.client.bucket(self.bucket_name)
#         blob = bucket.blob(blob_name)
#         try:
#             return blob.generate_signed_url(
#                 version="v4",
#                 expiration=timedelta(seconds=getattr(settings, "GCS_SIGNED_URL_TTL_SECONDS", 604800)),
#                 method="GET",
#                 response_disposition="inline",
#             )
#         except Exception:
#             return blob.public_url or f"[https://storage.googleapis.com/](https://storage.googleapis.com/){self.bucket_name}/{blob_name}"

#     def download_blob(self, blob_name: str) -> Optional[bytes]:
#         if not self.client or not self.bucket_name or not blob_name:
#             return None
#         try:
#             bucket = self.client.bucket(self.bucket_name)
#             blob = bucket.blob(blob_name)
#             return blob.download_as_bytes()
#         except Exception as exc:
#             return None


import mimetypes
import logging
from datetime import timedelta
from typing import Optional

from django.conf import settings
from google.cloud import storage
from google.api_core.retry import Retry
logger = logging.getLogger(__name__)

class GoogleStorageUploader:
    def __init__(self, bucket_name: Optional[str] = None):
        self.bucket_name = bucket_name or getattr(settings, "GCS_BUCKET_NAME", None)
        self.client = None
        
        # Only initialize client if we actually have a bucket to talk to
        if self.bucket_name:
            try:
                self.client = storage.Client()
            except Exception as e:
                logger.error(f"Failed to initialize GCS Client: {e}")

    def upload(self, name: str, content: bytes, content_type: Optional[str] = None) -> Optional[str]:
        if not self.client or not self.bucket_name:
            logger.error("GCS Upload skipped: No client or bucket configured.")
            return None
            
        bucket = self.client.bucket(self.bucket_name)
        blob = bucket.blob(name)
        
        # FIX: Auto-detect content type if missing. 
        # This ensures PDFs open in the browser ("inline") instead of downloading.
        if not content_type:
            content_type, _ = mimetypes.guess_type(name)
            
        try:
            blob.upload_from_string(
                content, 
                content_type=content_type, 
                timeout=120, 
                retry=Retry(deadline=180)
            )
            return blob.name
        except Exception as exc:
            logger.exception("GCS upload failed", extra={"blob": name, "error": str(exc)})
            return None

    def signed_url_for_blob(self, blob_name: str) -> Optional[str]:
        if not blob_name or not self.bucket_name:
            return None

        # Fallback URL (Cleaned up from your snippet)
        public_url = f"https://storage.googleapis.com/{self.bucket_name}/{blob_name}"

        if not self.client:
            return public_url

        try:
            bucket = self.client.bucket(self.bucket_name)
            blob = bucket.blob(blob_name)
            
            return blob.generate_signed_url(
                version="v4",
                expiration=timedelta(seconds=getattr(settings, "GCS_SIGNED_URL_TTL_SECONDS", 604800)), # 7 days
                method="GET",
                response_disposition="inline", # Tries to open in browser
            )
        except Exception as e:
            # Log the permission error so we know IF the Service Account is broken
            logger.warning(f"Failed to sign URL (Permissions issue?): {e}")
            return public_url

    def download_blob(self, blob_name: str) -> Optional[bytes]:
        if not self.client or not self.bucket_name or not blob_name:
            return None
        try:
            bucket = self.client.bucket(self.bucket_name)
            blob = bucket.blob(blob_name)
            return blob.download_as_bytes()
        except Exception as exc:
            logger.error(f"Failed to download blob {blob_name}: {exc}")
            return None


class M365InboxReader:
    def __init__(self, config=None):
        self.config = config or settings.M365_CONFIG

    def is_configured(self, log: bool = True) -> bool:
        required = ["tenant_id", "client_id", "client_secret"]
        missing = [key for key in required if not self.config.get(key)]
        if missing and log:
            return False
        return True

    def _client(self):
        if not self.is_configured(log=False):
            return None
        authority = f"[https://login.microsoftonline.com/](https://login.microsoftonline.com/){self.config.get('tenant_id')}"
        return msal.ConfidentialClientApplication(
            self.config.get("client_id"),
            authority=authority,
            client_credential=self.config.get("client_secret"),
        )

    def _fetch_attachments_for_message(self, message: dict, headers: dict, user: str) -> List[AttachmentPayload]:
        message_id = message["id"]
        sender = (message.get("from") or {}).get("emailAddress", {}).get("address")
        received_at = timezone.make_aware(datetime.strptime(message["receivedDateTime"], "%Y-%m-%dT%H:%M:%SZ"))
        attach_url = f"[https://graph.microsoft.com/v1.0/users/](https://graph.microsoft.com/v1.0/users/){user}/messages/{message_id}/attachments"
        try:
            attach_raw = requests.get(attach_url, headers=headers, timeout=30)
            attach_resp = attach_raw.json()
            if attach_raw.status_code >= 400:
                return []
        except Exception:
            return []

        found: List[AttachmentPayload] = []
        for attachment in attach_resp.get("value", []):
            if not attachment.get("contentBytes"):
                continue
            name = attachment.get("name", "attachment")
            if not name.lower().endswith((".pdf", ".txt", ".docx", ".doc")):
                continue
            try:
                content = base64.b64decode(attachment["contentBytes"])
            except Exception:
                continue
            found.append(
                AttachmentPayload(
                    message_id=message_id,
                    sender=sender,
                    attachment_name=name,
                    received_at=received_at,
                    content=content,
                    source=ResumeSource.EMAIL,
                )
            )
        return found

    def fetch_recent_attachments(self, limit: int = 5, fetch_all: bool = False) -> List[AttachmentPayload]:
        client = self._client()
        if not client:
            return []

        token = client.acquire_token_for_client(scopes=["[https://graph.microsoft.com/.default](https://graph.microsoft.com/.default)"])
        if "access_token" not in token:
            return []

        headers = {"Authorization": f"Bearer {token['access_token']}"}
        user = self.config.get("user_id", "me")
        top = min(limit, 100) if limit else 50
        base_url = f"[https://graph.microsoft.com/v1.0/users/](https://graph.microsoft.com/v1.0/users/){user}/messages"
        messages_url = f"{base_url}?$filter=hasAttachments eq true&$orderby=receivedDateTime desc&$top={top}"
        fallback_url = f"{base_url}?$orderby=receivedDateTime desc&$top={top}"
        messages = []
        next_url = messages_url
        attempted_fallback = False
        while next_url:
            resp_raw = requests.get(next_url, headers=headers, timeout=30)
            resp = resp_raw.json()
            if resp_raw.status_code >= 400:
                error_code = (resp.get("error") or {}).get("code")
                if not attempted_fallback and error_code == "InefficientFilter":
                    next_url = fallback_url
                    attempted_fallback = True
                    continue
                break
            messages.extend(resp.get("value", []))
            if not fetch_all and len(messages) >= limit:
                messages = messages[:limit]
                break
            next_url = resp.get("@odata.nextLink") if fetch_all else None

        attachments: List[AttachmentPayload] = []
        if not messages:
            return attachments

        max_workers = min(8, len(messages))
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_map = {
                executor.submit(self._fetch_attachments_for_message, message, headers, user): message.get("id")
                for message in messages
            }
            for future in as_completed(future_map):
                try:
                    attachments.extend(future.result() or [])
                except Exception:
                    pass
        return attachments



# class ResumeIngestor:
#     def __init__(
#         self,
#         scorer: Optional['GeminiEvaluator'] = None,
#         uploader: Optional['GoogleStorageUploader'] = None,
#         max_workers: Optional[int] = None,
#         score_workers: Optional[int] = None,
#     ):
#         # REMOVED the invalid import: from .ai import GeminiEvaluator
        
#         self.scorer = scorer
#         # If GoogleStorageUploader is in this file, use it directly. 
#         # If it's in a separate file (e.g. storage.py), keep the import but fix the path.
#         # Assuming for now it is in services.py or imported at top:
#         self.uploader = uploader or GoogleStorageUploader() 
        
#         self.max_workers = max_workers or getattr(settings, "INGEST_MAX_WORKERS", 4)
#         self.score_workers = score_workers or getattr(settings, "INGEST_SCORE_WORKERS", self.max_workers)
#         self._scorer_local = threading.local()

#     def _get_scorer(self) -> 'GeminiEvaluator':
#         # REMOVED the invalid import here too.
#         if not getattr(self._scorer_local, "instance", None):
#             # Just instantiate it directly since it's in the same file
#             self._scorer_local.instance = self.scorer or GeminiEvaluator()
#         return self._scorer_local.instance

#     def _score_resume_for_jobs(self, resume: Resume, text_content: str, jobs: List[JobDescription]) -> None:
#         if not jobs:
#             return
        
#         max_workers = max(1, min(self.score_workers, len(jobs)))

#         def score_job(job: JobDescription):
#             close_old_connections()
#             scorer = self._get_scorer()
#             try:
#                 total, detail = scorer.score_resume_against_job(text_content, job)
#                 return job, total, detail
#             finally:
#                 close_old_connections()

#         results = []
#         with ThreadPoolExecutor(max_workers=max_workers) as executor:
#             future_map = {executor.submit(score_job, job): job for job in jobs}
#             for future in as_completed(future_map):
#                 try:
#                     job, total, detail = future.result()
#                     results.append((job, total, detail))
#                 except Exception as exc:
#                     job_ref = future_map[future]
#                     logger.exception("Scoring failed", extra={"job_id": job_ref.id, "error": str(exc)})

#         for job, total, detail in results:
#             # 1. Create/Update the score object
#             score_obj, created = ResumeScore.objects.update_or_create(
#                 resume=resume,
#                 job=job,
#                 defaults={"total_score": total, "detail": detail},
#             )
            
#             # 2. TRIGGER AUTOMATION (Assuming check_and_process_automation is defined in this file)
#             try:
#                 check_and_process_automation(score_obj)
#             except Exception as e:
#                 logger.error(f"Automation trigger failed for resume {resume.id}: {e}")

#     def _record_failure(self, attachment: AttachmentPayload, jobs: List[JobDescription], error: Exception) -> None:
#         try:
#             sender_clean = clean_string(attachment.sender or "") if attachment.sender else None
#             candidate_name = clean_string(attachment.attachment_name.rsplit(".", 1)[0])[:255]
            
#             resume, _ = Resume.objects.get_or_create(
#                 message_id=attachment.message_id,
#                 attachment_name=attachment.attachment_name,
#                 defaults={
#                     "sender": sender_clean,
#                     "candidate_email": sender_clean,
#                     "candidate_name": candidate_name,
#                     "text_content": "",
#                     "received_at": attachment.received_at,
#                     "source": attachment.source or ResumeSource.EMAIL,
#                 },
#             )
            
#             detail = [{"title": "Processing failed", "score": 0, "notes": truncate(str(error), 200)}]
#             for job in jobs:
#                 ResumeScore.objects.update_or_create(
#                     resume=resume,
#                     job=job,
#                     defaults={"total_score": 0.0, "detail": detail},
#                 )
#         except Exception:
#             pass

#     def _ingest_single(self, attachment: AttachmentPayload, jobs: List[JobDescription]) -> int:
#         close_old_connections()
#         try:
#             raw_text = extract_text_from_attachment(attachment.attachment_name, attachment.content)
#             text_content = clean_string(raw_text)
#             phone = extract_phone_number(text_content)
#             text_email = extract_email_from_text(text_content)
            
#             sender_clean = clean_string(attachment.sender or "") if attachment.sender else None
#             candidate_email = text_email or sender_clean

#             storage_path = f"resumes/{attachment.message_id}/{attachment.attachment_name}"
#             content_type = "application/pdf" if attachment.attachment_name.lower().endswith(".pdf") else "text/plain"
            
#             file_url = truncate(
#                 self.uploader.upload(storage_path, attachment.content, content_type=content_type) or "",
#                 2000,
#             )
#             candidate_name = clean_string(attachment.attachment_name.rsplit(".", 1)[0])[:255]

#             # CHECK EXISTING
#             existing_resume = None
#             if candidate_email or phone:
#                 query = Q()
#                 if candidate_email:
#                     query |= Q(candidate_email=candidate_email) | Q(sender=candidate_email)
#                 if phone:
#                     query |= Q(candidate_phone=phone)
#                 existing_resume = Resume.objects.filter(query).order_by('-received_at').first()

#             # CREATE OR UPDATE
#             if existing_resume:
#                 logger.info(f"Updating existing resume for {candidate_email or phone}")
#                 resume = existing_resume
#                 resume.file_url = file_url
#                 resume.text_content = text_content
#                 resume.received_at = attachment.received_at
#                 resume.message_id = attachment.message_id
#                 resume.attachment_name = attachment.attachment_name
                
#                 if candidate_email: resume.candidate_email = candidate_email
#                 if phone: resume.candidate_phone = phone
#                 if candidate_name: resume.candidate_name = candidate_name

#                 resume.status = 'PENDING'
#                 resume.human_score = None
#                 resume.human_score_breakdown = {}
#                 resume.is_active_seeker = infer_active_seeker(sender_clean, candidate_email, attachment.source)
#                 resume.save()
#                 log_activity(resume, 'INGESTED', f"Updated via {attachment.source or 'Unknown'}")
#             else:
#                 logger.info(f"Creating new resume for {candidate_email or phone}")
#                 resume = Resume.objects.create(
#                     message_id=attachment.message_id,
#                     attachment_name=attachment.attachment_name,
#                     sender=sender_clean,
#                     candidate_email=candidate_email,
#                     candidate_phone=phone,
#                     candidate_name=candidate_name,
#                     text_content=text_content,
#                     file_url=file_url or "",
#                     received_at=attachment.received_at,
#                     source=attachment.source or ResumeSource.EMAIL,
#                     is_active_seeker=infer_active_seeker(sender_clean, candidate_email, attachment.source),
#                     status='PENDING'
#                 )
#                 log_activity(resume, 'INGESTED', f"New application via {attachment.source or 'Unknown'}")

#             self._score_resume_for_jobs(resume, text_content, jobs)
#             return 1

#         except Exception as exc:
#             logger.exception("Resume ingest failed", extra={"error": str(exc)})
#             self._record_failure(attachment, jobs, exc)
#             return 0
#         finally:
#             close_old_connections()

#     def ingest(self, attachments: List[AttachmentPayload], jobs: Optional[List[JobDescription]] = None) -> int:
#         if not attachments:
#             return 0

#         jobs = jobs or list(JobDescription.objects.filter(active=True))
#         max_workers = max(1, min(self.max_workers, len(attachments)))
#         saved_count = 0
        
#         with ThreadPoolExecutor(max_workers=max_workers) as executor:
#             futures = [executor.submit(self._ingest_single, attachment, jobs) for attachment in attachments]
#             for future in as_completed(futures):
#                 try:
#                     saved_count += future.result() or 0
#                 except Exception:
#                     pass
#         return saved_count


import logging
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Optional
from django.conf import settings
from django.db.models import Q
from django.db import close_old_connections

# Assumes these models/functions are imported or defined above
# from .models import Resume, JobDescription, ResumeScore, AttachmentPayload, ResumeSource
# from .utils import clean_string, truncate, extract_text_from_attachment, extract_phone_number, extract_email_from_text, infer_active_seeker, log_activity

logger = logging.getLogger(__name__)

class ResumeIngestor:
    def __init__(
        self,
        scorer: Optional['GeminiEvaluator'] = None,
        uploader: Optional['GoogleStorageUploader'] = None,
        max_workers: Optional[int] = None,
        score_workers: Optional[int] = None,
    ):
        self.scorer = scorer
        # Use the class directly since it's in the same file/module
        self.uploader = uploader or GoogleStorageUploader() 
        
        self.max_workers = max_workers or getattr(settings, "INGEST_MAX_WORKERS", 4)
        self.score_workers = score_workers or getattr(settings, "INGEST_SCORE_WORKERS", self.max_workers)
        self._scorer_local = threading.local()

    def _get_scorer(self) -> 'GeminiEvaluator':
        """Thread-safe access to GeminiEvaluator"""
        if not getattr(self._scorer_local, "instance", None):
            self._scorer_local.instance = self.scorer or GeminiEvaluator()
        return self._scorer_local.instance

    # def _score_resume_for_jobs(self, resume: Resume, text_content: str, jobs: List[JobDescription]) -> None:
    #     if not jobs:
    #         return
        
    #     # Limit concurrency to avoid hitting API limits
    #     max_workers = max(1, min(self.score_workers, len(jobs)))

    #     def score_job(job: JobDescription):
    #         close_old_connections() # Safety for DB in threads
    #         scorer = self._get_scorer()
    #         try:
    #             # FIX: Unpack 3 values (Total, Detail, Profile)
    #             total, detail, profile = scorer.score_resume_against_job(text_content, job)
    #             return job, total, detail, profile
    #         finally:
    #             close_old_connections()

    #     results = []
    #     with ThreadPoolExecutor(max_workers=max_workers) as executor:
    #         future_map = {executor.submit(score_job, job): job for job in jobs}
    #         for future in as_completed(future_map):
    #             try:
    #                 job, total, detail, profile = future.result()
    #                 results.append((job, total, detail, profile))
    #             except Exception as exc:
    #                 job_ref = future_map[future]
    #                 logger.exception("Scoring failed", extra={"job_id": job_ref.id, "error": str(exc)})

    #     # Process results
    #     for job, total, detail, profile in results:
    #         # 1. Update Candidate Metadata (if Gemini found better info than Regex)
    #         if profile:
    #             updated = False
    #             if profile.get("name") and not resume.candidate_name:
    #                 resume.candidate_name = profile["name"]
    #                 updated = True
    #             if profile.get("email") and not resume.candidate_email:
    #                 resume.candidate_email = profile["email"]
    #                 updated = True
    #             if profile.get("linkedin_url") and not resume.linkedin_profile:
    #                 resume.linkedin_profile = profile["linkedin_url"]
    #                 updated = True
    #             if profile.get("phone") and not resume.candidate_phone:
    #                 resume.candidate_phone = profile["phone"]
    #                 updated = True
                
    #             if updated:
    #                 resume.save(update_fields=["candidate_name", "candidate_email", "linkedin_profile", "candidate_phone"])

    #         # 2. Create/Update Score
    #         score_obj, created = ResumeScore.objects.update_or_create(
    #             resume=resume,
    #             job=job,
    #             defaults={"total_score": total, "detail": detail},
    #         )
            
    #         # 3. Trigger Automation (Safe Call)
    #         if "check_and_process_automation" in globals():
    #             try:
    #                 check_and_process_automation(score_obj)
    #             except Exception as e:
    #                 logger.error(f"Automation trigger failed for resume {resume.id}: {e}")

    def _score_resume_for_jobs(self, resume: Resume, text_content: str, jobs: List[JobDescription]) -> None:
        if not jobs:
            return
        
        # Limit concurrency to avoid hitting API limits
        max_workers = max(1, min(self.score_workers, len(jobs)))

        def score_job(job: JobDescription):
            close_old_connections()
            scorer = self._get_scorer()
            try:
                # Unpack 3 values
                total, detail, profile = scorer.score_resume_against_job(text_content, job)
                return job, total, detail, profile
            finally:
                close_old_connections()

        results = []
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_map = {executor.submit(score_job, job): job for job in jobs}
            for future in as_completed(future_map):
                try:
                    job, total, detail, profile = future.result()
                    results.append((job, total, detail, profile))
                except Exception as exc:
                    job_ref = future_map[future]
                    logger.exception("Scoring failed", extra={"job_id": job_ref.id, "error": str(exc)})

        # Process results
        for job, total, detail, profile in results:
            
            # 1. Update Candidate Metadata
            if profile:
                updated_fields = []
                
                # Helper to update and track fields
                def update_if_new(field, value):
                    if value and not getattr(resume, field):
                        setattr(resume, field, value)
                        updated_fields.append(field)

                # Standard Fields
                update_if_new("candidate_name", profile.get("name"))
                update_if_new("candidate_email", profile.get("email"))
                update_if_new("candidate_phone", profile.get("phone"))
                update_if_new("linkedin_profile", profile.get("linkedin_url"))
                
                # NEW Fields (Role, Company, Location)
                update_if_new("candidate_role", profile.get("current_role"))
                update_if_new("candidate_company", profile.get("current_company"))
                update_if_new("candidate_location", profile.get("current_location"))
                
                if updated_fields:
                    resume.save(update_fields=updated_fields)

            # 2. Create/Update Score
            score_obj, created = ResumeScore.objects.update_or_create(
                resume=resume,
                job=job,
                defaults={"total_score": total, "detail": detail},
            )
            
            # 3. Trigger Automation (Use explicit import)
            try:
                # Import here to avoid circular dependency
                from .services import check_and_process_automation
                check_and_process_automation(score_obj)
            except Exception as e:
                logger.error(f"Automation trigger failed for resume {resume.id}: {e}")

    # def _ingest_single(self, attachment: AttachmentPayload, jobs: List[JobDescription]) -> int:
    #     close_old_connections()
    #     try:
    #         # 1. Extract Text
    #         raw_text = extract_text_from_attachment(attachment.attachment_name, attachment.content)
    #         text_content = clean_string(raw_text)
            
    #         # 2. Initial Regex Extraction (Fallback)
    #         phone = extract_phone_number(text_content)
    #         text_email = extract_email_from_text(text_content)
            
    #         sender_clean = clean_string(attachment.sender or "") if attachment.sender else None
    #         candidate_email = text_email or sender_clean

    #         # 3. Upload File
    #         storage_path = f"resumes/{attachment.message_id}/{attachment.attachment_name}"
    #         # Use mimetypes guessing logic inside uploader, or fallback here
    #         content_type = "application/pdf" if attachment.attachment_name.lower().endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
    #         file_url = truncate(
    #             self.uploader.upload(storage_path, attachment.content, content_type=content_type) or "",
    #             2000,
    #         )
            
    #         # Simple name guess (Gemini will correct this later)
    #         candidate_name = clean_string(attachment.attachment_name.rsplit(".", 1)[0])[:255]

    #         # 4. Check Existing Resume
    #         existing_resume = None
    #         if candidate_email or phone:
    #             query = Q()
    #             if candidate_email:
    #                 query |= Q(candidate_email=candidate_email) | Q(sender=candidate_email)
    #             if phone:
    #                 query |= Q(candidate_phone=phone)
    #             existing_resume = Resume.objects.filter(query).order_by('-received_at').first()

    #         # 5. Save to DB
    #         if existing_resume:
    #             logger.info(f"Updating existing resume for {candidate_email or phone}")
    #             resume = existing_resume
    #             resume.file_url = file_url
    #             resume.text_content = text_content
    #             resume.received_at = attachment.received_at
    #             resume.message_id = attachment.message_id
    #             resume.attachment_name = attachment.attachment_name
                
    #             # Only update basic fields if missing, let Gemini fill the rest
    #             if candidate_email and not resume.candidate_email: resume.candidate_email = candidate_email
    #             if phone and not resume.candidate_phone: resume.candidate_phone = phone

    #             resume.status = 'PENDING'
    #             resume.is_active_seeker = infer_active_seeker(sender_clean, candidate_email, attachment.source)
    #             resume.save()
    #             log_activity(resume, 'INGESTED', f"Updated via {attachment.source or 'Unknown'}")
    #         else:
    #             logger.info(f"Creating new resume for {candidate_email or phone}")
    #             resume = Resume.objects.create(
    #                 message_id=attachment.message_id,
    #                 attachment_name=attachment.attachment_name,
    #                 sender=sender_clean,
    #                 candidate_email=candidate_email,
    #                 candidate_phone=phone,
    #                 candidate_name=candidate_name,
    #                 text_content=text_content,
    #                 file_url=file_url or "",
    #                 received_at=attachment.received_at,
    #                 source=attachment.source or ResumeSource.EMAIL,
    #                 is_active_seeker=infer_active_seeker(sender_clean, candidate_email, attachment.source),
    #                 status='PENDING'
    #             )
    #             log_activity(resume, 'INGESTED', f"New application via {attachment.source or 'Unknown'}")

    #         # 6. Score & Enrich (This calls Gemini)
    #         self._score_resume_for_jobs(resume, text_content, jobs)
    #         return 1

    #     except Exception as exc:
    #         logger.exception("Resume ingest failed", extra={"error": str(exc)})
    #         # Ensure _record_failure is defined in your file
    #         # self._record_failure(attachment, jobs, exc) 
    #         return 0
    #     finally:
    #         close_old_connections()

    def _ingest_single(self, attachment: AttachmentPayload, jobs: List[JobDescription]) -> int:
        close_old_connections()
        try:
            # 1. Extract Text (Use the NEW robust extractor)
            # Ensure extract_text_with_links is imported from your utils/services
            try:
                from .services import extract_text_with_links
                # raw_text = extract_text_with_links(attachment.content) # Pass stream/bytes
                raw_text = extract_text_with_links(attachment.content, filename=attachment.attachment_name)
            except ImportError:
                # Fallback if you haven't moved the function yet
                raw_text = extract_text_from_attachment(attachment.attachment_name, attachment.content)
            
            text_content = clean_string(raw_text)
            
            # 2. Initial Regex Extraction (Fallback)
            phone = extract_phone_number(text_content)
            text_email = extract_email_from_text(text_content)
            
            sender_clean = clean_string(attachment.sender or "") if attachment.sender else None
            candidate_email = text_email or sender_clean

            # 3. Upload File
            storage_path = f"resumes/{attachment.message_id}/{attachment.attachment_name}"
            content_type = "application/pdf" if attachment.attachment_name.lower().endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            file_url = truncate(
                self.uploader.upload(storage_path, attachment.content, content_type=content_type) or "",
                2000,
            )
            
            candidate_name = clean_string(attachment.attachment_name.rsplit(".", 1)[0])[:255]

            # 4. Check Existing Resume
            existing_resume = None
            if candidate_email or phone:
                query = Q()
                if candidate_email:
                    query |= Q(candidate_email=candidate_email) | Q(sender=candidate_email)
                if phone:
                    query |= Q(candidate_phone=phone)
                existing_resume = Resume.objects.filter(query).order_by('-received_at').first()

            # 5. Save to DB
            if existing_resume:
                logger.info(f"Updating existing resume for {candidate_email or phone}")
                resume = existing_resume
                resume.file_url = file_url
                resume.text_content = text_content
                resume.received_at = attachment.received_at
                resume.message_id = attachment.message_id
                resume.attachment_name = attachment.attachment_name
                
                if candidate_email and not resume.candidate_email: resume.candidate_email = candidate_email
                if phone and not resume.candidate_phone: resume.candidate_phone = phone

                resume.status = 'PENDING'
                resume.save()
                log_activity(resume, 'INGESTED', f"Updated via {attachment.source or 'Unknown'}")
            else:
                logger.info(f"Creating new resume for {candidate_email or phone}")
                resume = Resume.objects.create(
                    message_id=attachment.message_id,
                    attachment_name=attachment.attachment_name,
                    sender=sender_clean,
                    candidate_email=candidate_email,
                    candidate_phone=phone,
                    candidate_name=candidate_name,
                    text_content=text_content,
                    file_url=file_url or "",
                    received_at=attachment.received_at,
                    source=attachment.source or ResumeSource.EMAIL,
                    is_active_seeker=infer_active_seeker(sender_clean, candidate_email, attachment.source),
                    status='PENDING'
                )
                log_activity(resume, 'INGESTED', f"New application via {attachment.source or 'Unknown'}")

            # 6. Score & Enrich
            self._score_resume_for_jobs(resume, text_content, jobs)
            return 1

        except Exception as exc:
            logger.exception("Resume ingest failed", extra={"error": str(exc)})
            return 0
        finally:
            close_old_connections()

    def ingest(self, attachments: List[AttachmentPayload], jobs: Optional[List[JobDescription]] = None) -> int:
        if not attachments:
            return 0

        jobs = jobs or list(JobDescription.objects.filter(active=True))
        max_workers = max(1, min(self.max_workers, len(attachments)))
        saved_count = 0
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = [executor.submit(self._ingest_single, attachment, jobs) for attachment in attachments]
            for future in as_completed(futures):
                try:
                    saved_count += future.result() or 0
                except Exception:
                    pass
        return saved_count


# def populate_resume_text(resume: Resume, uploader: Optional[GoogleStorageUploader] = None) -> bool:
#     if resume.text_content:
#         return False
#     uploader = uploader or GoogleStorageUploader()
#     content_bytes: Optional[bytes] = None
#     try:
#         if resume.file_url:
#             if resume.file_url.startswith("http"):
#                 resp = requests.get(resume.file_url, timeout=20)
#                 if resp.status_code < 400:
#                     content_bytes = resp.content
#             else:
#                 content_bytes = uploader.download_blob(resume.file_url)
#     except Exception as exc:
#         record_error("populate_resume_text", str(exc), resume=resume)
#         return False

#     if not content_bytes:
#         return False
#     text_content = clean_string(extract_text_from_attachment(resume.attachment_name, content_bytes))
#     if not text_content:
#         return False
#     resume.text_content = text_content
#     resume.save(update_fields=["text_content"])
#     return True

def populate_resume_text(resume: Resume, uploader: Optional[GoogleStorageUploader] = None) -> bool:
    if resume.text_content:
        return False

    uploader = uploader or GoogleStorageUploader()
    content_bytes: Optional[bytes] = None
    
    try:
        if resume.file_url:
            # Handle full HTTP URLs (Legacy data or external uploads)
            if resume.file_url.startswith("http"):
                resp = requests.get(resume.file_url, timeout=20)
                if resp.status_code < 400:
                    content_bytes = resp.content
            # Handle GCS Blob Names (New Standard)
            else:
                content_bytes = uploader.download_blob(resume.file_url)

        if not content_bytes:
            # Last resort: Try downloading by attachment name if path is missing
            # Construct path: resumes/{message_id}/{attachment_name}
            fallback_path = f"resumes/{resume.message_id}/{resume.attachment_name}"
            content_bytes = uploader.download_blob(fallback_path)

    except Exception as exc:
        record_error("populate_resume_text", str(exc), resume=resume)
        return False

    if not content_bytes:
        return False

    # Extract text using our new robust function
    text_content = clean_string(extract_text_from_attachment(resume.attachment_name, content_bytes))
    
    if not text_content:
        return False

    resume.text_content = text_content
    resume.save(update_fields=["text_content"])
    return True


# jobs/automation.py

# def check_and_process_automation(resume_score):
#     resume = resume_score.resume
#     score_val = resume_score.total_score
#     threshold = getattr(settings, 'AUTO_REJECTION_THRESHOLD', 60) 

#     # 1. AUTO REJECTION LOGIC
#     if score_val < threshold:
#         # Only update and log if it's a *new* rejection
#         if resume.status != 'AUTO_REJECTED':
#             resume.status = 'AUTO_REJECTED'
#             resume.save()
#             log_activity(resume, 'STATUS_CHANGE', f"Auto-rejected: Score {score_val} < {threshold}")

#         # --- EMAIL SENDING DISABLED (COMMENTED OUT) ---
#         # if not resume.rejection_email_sent and resume.candidate_email:
#         #     subject = f"Update regarding your application for {resume_score.job.name}"
#         #     message = (
#         #         f"Dear {resume.candidate_name or 'Candidate'},\n\n"
#         #         f"Thank you for your interest in the {resume_score.job.name} position. "
#         #         "We appreciate the time you took to apply.\n\n"
#         #         "After reviewing your application, we have decided not to proceed with your candidacy "
#         #         "at this time. We will keep your resume on file for future openings.\n\n"
#         #         "Best regards,\nHR Team"
#         #     )
#         #     try:
#         #         send_mail(subject, message, settings.DEFAULT_FROM_EMAIL, [resume.candidate_email], fail_silently=False)
#         #         
#         #         resume.rejection_email_sent = True
#         #         resume.save()
#         #         
#         #         log_activity(resume, 'EMAIL_SENT', "Auto-rejection email sent to candidate.")
#         #         
#         #     except Exception as e:
#         #         record_error("automation:email", str(e))
#         # ----------------------------------------------

#     # 2. SUCCESS LOGIC (Qualified)
#     else:
#         # Only move to NEW if it is currently PENDING.
#         if resume.status == 'PENDING':
#             resume.status = 'NEW'
#             resume.save()
#             log_activity(resume, 'STATUS_CHANGE', f"Qualified as NEW (Score: {score_val})")

# jobs/automation.py

def check_and_process_automation(resume_score):
    resume = resume_score.resume
    score_val = resume_score.total_score
    threshold = getattr(settings, 'AUTO_REJECTION_THRESHOLD', 60) 

    # 1. AUTO REJECTION LOGIC
    if score_val < threshold:
        # Only reject if not already rejected or hired
        if resume.status not in ['AUTO_REJECTED', 'REJECTED', 'HIRED', 'INTERVIEW_SCHEDULED']:
            resume.status = 'AUTO_REJECTED'
            resume.save()
            log_activity(resume, 'STATUS_CHANGE', f"Auto-rejected: Score {score_val} < {threshold}")

            # Note: Email sending logic is correctly commented out as per your request.
            # --- EMAIL SENDING DISABLED (COMMENTED OUT) ---
            # if not resume.rejection_email_sent and resume.candidate_email:
            #     subject = f"Update regarding your application for {resume_score.job.name}"
            #     message = (
            #         f"Dear {resume.candidate_name or 'Candidate'},\n\n"
            #         f"Thank you for your interest in the {resume_score.job.name} position. "
            #         "We appreciate the time you took to apply.\n\n"
            #         "After reviewing your application, we have decided not to proceed with your candidacy "
            #         "at this time. We will keep your resume on file for future openings.\n\n"
            #         "Best regards,\nHR Team"
            #     )
            #     try:
            #         send_mail(subject, message, settings.DEFAULT_FROM_EMAIL, [resume.candidate_email], fail_silently=False)
            #         
            #         resume.rejection_email_sent = True
            #         resume.save()
            #         
            #         log_activity(resume, 'EMAIL_SENT', "Auto-rejection email sent to candidate.")
            #         
            #     except Exception as e:
            #         record_error("automation:email", str(e))
            # ----------------------------------------------


    # 2. SUCCESS LOGIC (Qualified)
    else:
        # If score improves (e.g., after rescoring), we might want to "Un-reject" them
        # Only move to NEW if they were PENDING or AUTO_REJECTED (giving them a second chance)
        if resume.status in ['PENDING', 'AUTO_REJECTED']:
            resume.status = 'NEW'
            resume.save()
            log_activity(resume, 'STATUS_CHANGE', f"Qualified as NEW (Score: {score_val} >= {threshold})")